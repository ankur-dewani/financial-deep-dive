"""
DD1 Document Generator
Fills the Deep Dive 1 template with analysis results.

Template structure (single table, 8 rows x 2 columns):
  Row 0: Function         -> "Operations"  (pre-filled)
  Row 1: Playbook Item    -> "Initial Import" (pre-filled)
  Row 2: Problem Statement
  Row 3: 5 Why Analysis
  Row 4: Root Cause
  Row 5: Fix
  Row 6: AI Opportunities
  Row 7: AI Tools Used

Writing style rules:
  - No em dashes, en dashes, or hyphens between words
  - No AI language ("comprehensive", "leverage", "high-impact")
  - Professional, strategic, simple English
  - CEO reads this personally
"""

from __future__ import annotations

import os
import sys
import shutil

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from cost_model import get_savings_summary, get_target_fa_model, get_current_fa_cost

# ---------------------------------------------------------------------------
# File paths
# ---------------------------------------------------------------------------
TEMPLATE = os.path.join(
    os.path.dirname(__file__), "..", "data", "input",
    "Deep Dive 1 - Template.docx"
)
OUTPUT = os.path.join(
    os.path.dirname(__file__), "..", "data", "output",
    "DD1 - Ankur Dewani.docx"
)


# ---------------------------------------------------------------------------
# Content for each row
# ---------------------------------------------------------------------------

PROBLEM_STATEMENT = (
    "What:\n"
    "This business unit generates $79.2M in revenue but retains only $7M "
    "(8.86% margin) versus the 70% target. Expenses consume 91.1% of revenue "
    "against a 30% benchmark.\n\n"
    "Why:\n"
    "The acquisition was never operationally transformed to fit the Central Factory model. "
    "Every functional area exceeds its benchmark: Shared Services runs at 12% vs 4.5%, "
    "Engineering at 21.8% vs 10%, Sales at 14.4% vs 5%, and Product at 17.2% vs 2%. "
    "No function has been restructured around standardized roles or centralized delivery."
)

FIVE_WHY_ANALYSIS = (
    "Question 1: Why is this business unit not in model?\n"
    "Answer: The business unit still runs its pre-acquisition cost structure. No function was restructured "
    "for the Central Factory model. Expenses total $72.2M (91.1% of revenue vs 30% benchmark) with all 8 "
    "benchmark categories over target. The largest absolute gaps: Engineering +$9.4M, Hosting +$6.6M, "
    "Product +$12.0M, Sales +$7.4M.\n"
    "Evidence: Benchmark Mapping sheet. Every category shows 'Over' status.\n\n"

    "Question 2: Why start with Shared Services when Engineering has a larger gap?\n"
    "Answer: Engineering ($9.4M gap) requires building a target role structure from scratch. Shared Services "
    "has a pre-defined Central Finance model with exact roles and rates, enabling a specific fix with quantified "
    "savings in 16 weeks. This proves the integration methodology works before applying it to larger, more complex "
    "functions. Within Shared Services ($15.8M, 12% of revenue, 74 employees, 11 G&A sub-departments), F&A alone "
    "costs $3.82M (4.8% of revenue) and exceeds the entire 4.5% Shared Services benchmark by itself.\n"
    "Evidence: SS Breakdown sheet. F&A highlighted: 18 employees, $1.45M headcount, $2.37M non-HC OPEX.\n\n"

    "Question 3: Why does Finance and Accounting cost $3.82M?\n"
    "Answer: 51% of F&A spend ($1.96M) flows to external providers. $1.43M goes to outsourced services "
    "(audit firms, tax advisors, bookkeeping) and $529K to external contractors. "
    "The 18 internal employees ($1.45M) spend their time coordinating vendor handoffs rather than executing "
    "standardized finance processes directly.\n"
    "Evidence: FA Deep Dive sheet. Outsourced Services ($1,426,248) and External Contractors ($529,469) "
    "together represent 51.2% of total F&A cost.\n\n"

    "Question 4: Why is 51% of F&A spend going to external providers?\n"
    "Answer: Each acquired entity retained its own audit firm, tax advisor, and accounting processes. "
    "No consolidation happened. The result: duplicate audit engagements across entities, "
    "no unified chart of accounts, no shared financial platform, and no standardized close process. "
    "Each entity runs an independent close cycle with its own vendor stack.\n"
    "Evidence: OPEX line items show multiple audit and accounting vendors. "
    "Outsourced Services is the single largest F&A cost category.\n\n"

    "Question 5: Why do these fragmented processes persist?\n"
    "Answer: No one was tasked with operational integration. The acquisition brought revenue into the portfolio "
    "but left the cost structure untouched. Without a mandate to simplify processes and transition to Central Finance "
    "roles (VP, Manager, Senior Accountant, Accountant), each entity continued operating independently. "
    "The absence of a Central Finance function is a leadership gap, not a technical one.\n"
    "Evidence: F&A employee salaries range from $35K to $176K with no alignment to Central Finance role tiers. "
    "FA Employee Analysis sheet shows the mapping gap."
)

ROOT_CAUSE = (
    "No one was tasked with integrating Finance and Accounting into the Central Finance model.\n\n"
    "The function operates as a fragmented multi-entity operation where 51% of cost ($1.96M) flows to "
    "external providers doing duplicated work across entities, and 18 internal staff coordinate vendor "
    "handoffs rather than executing standardized processes.\n\n"
    "At $3.82M (4.8% of revenue), F&A alone exceeds the entire Shared Services benchmark of 4.5%. "
    "This is a leadership gap: no unified chart of accounts, "
    "no consolidated audit engagement, no shared financial platform, and no standardized role structure. "
    "The fix exists. The Central Finance Roles define the target. The gap is execution."
)


def _build_fix_text() -> str:
    """Build the Fix section text using cost model data."""
    target = get_target_fa_model()
    summary = get_savings_summary()
    current = get_current_fa_cost()

    lines = [
        "Simplify Finance and Accounting first, then move to the Central Finance model.\n",
        "Customer impact: Zero. F&A is a back-office function. No customer facing "
        "processes change. The 87% recurring revenue base remains unaffected.\n",
        "Target Central Finance Team (18 roles, standardized):\n",
    ]

    # Role table
    for role in target["roles"]:
        lines.append(
            f"  {role['count']}x {role['role']}: ${role['count'] * role['annual']:,}/year"
        )

    lines.append(f"  Statutory audit (retained): ${target['statutory_audit']:,}/year")
    lines.append(f"  Total in-model cost: ${target['total']:,}/year\n")

    lines.append("Cost Bridge:")
    lines.append(f"  Current F&A cost: ${summary['current_total']:,}")
    lines.append(f"  Target in-model cost: ${summary['target_total']:,}")
    lines.append(f"  Annual savings: ${summary['annual_savings']:,} ({summary['savings_pct']:.0%} reduction)")
    lines.append(f"  Margin impact: +{summary['annual_savings']/79_194_484:.1%} of revenue\n")

    lines.append("Why 18 people can absorb the outsourced work:")
    lines.append("  Consolidating 7+ audit engagements into 1 or 2 eliminates duplicate entity-level work.")
    lines.append("  One unified chart of accounts replaces multiple entity-specific structures.")
    lines.append("  One close process replaces parallel close cycles per entity.")
    lines.append("  18 staff currently spend time on vendor coordination that disappears when work is in house.")
    lines.append("  AI close automation (see AI Opportunities) absorbs remaining volume increase.\n")

    lines.append("Feasibility: HIGH")
    lines.append("  Central Finance Roles spreadsheet provides exact target costs.")
    lines.append("  F&A is back-office. Zero customer facing processes change.")
    lines.append("  This is the first function to prove the integration methodology. Engineering and Sales follow.\n")

    lines.append("Phased Implementation:\n")
    lines.append("Phase 1: Simplify and document (weeks 1 to 4)")
    lines.append("  Freeze new outsourced engagements immediately.")
    lines.append("  Document every F&A process across all entities. Map duplicates.")
    lines.append("  Map each of 18 employees to a Central Finance role tier by salary band and skills.")
    lines.append("  Write the standardized close checklist: 5 day close, single GL, automated intercompany eliminations.")
    lines.append("  KPI: 100% of F&A processes documented, simplified, ready for Central.\n")

    lines.append("Phase 2: Migrate to Central (weeks 5 to 12)")
    lines.append("  Extend existing Central finance platform to this BU (or configure NetSuite if no Central tool exists).")
    lines.append("  Migrate chart of accounts to unified structure.")
    lines.append("  Consolidate audit firms from 7+ to 1 or 2.")
    lines.append("  Bring outsourced work in house using simplified processes from Phase 1.")
    lines.append("  KPI: Reduce outsourced spend by 75% ($1.47M).\n")

    lines.append("Phase 3: Optimization (weeks 13 to 16)")
    lines.append("  Run first full close cycle on Central Finance model.")
    lines.append("  Terminate remaining vendor contracts.")
    lines.append("  Configure AI close automation on unified GL (see AI Opportunities).")
    lines.append("  KPI: F&A cost under $1.2M annualized run rate.\n")

    lines.append("Risks:")
    lines.append("  Statutory compliance: Retain external audit for legal requirements ($200K).")
    lines.append("  Employee transition: Employees who do not fit Central Finance roles exit. Budget 2 to 3 replacements.")
    lines.append("  Multi-jurisdiction tax: Retain 1 to 2 outsourced tax advisors for non-US filings during transition. Budget $50K to $75K.")

    return "\n".join(lines)


AI_OPPORTUNITIES = (
    "Area 1: LLM Driven Financial Close\n\n"
    "Process: Monthly close requires manual intercompany reconciliation, journal entry preparation, "
    "and exception review across multiple entities.\n"
    "Current cost driver: Close cycle takes 15 to 20 days. Accountants spend time on matching "
    "and data entry rather than analysis.\n"
    "AI solution: I configure an LLM reconciliation agent on the unified GL (post migration). "
    "It auto-matches 100% of intercompany transactions, generates recurring journal entries, "
    "and routes only true exceptions for human review. This replaces manual matching entirely.\n"
    "Estimated impact: Close cycle drops from 15 to 20 days to under 5 days. "
    "Reduces Senior Accountant need from 5 to 3, saving an additional $120K/year on top of the structural fix.\n"
    "Complexity: Medium. Requires unified GL from Phase 2 and 8 to 12 weeks configuration.\n\n"
    "Area 2: AI Spend Intelligence\n\n"
    "Process: Vendor spend classification and duplicate payment detection are currently manual quarterly exercises.\n"
    "Current cost driver: $1.96M in outsourced F&A spend had no real time monitoring. "
    "Cost creep went undetected between review cycles.\n"
    "AI solution: I build a Claude Code pipeline on existing AP data feeds. It classifies vendor spend daily, "
    "detects duplicates, flags benchmark non-compliance, and produces a weekly list of vendors to cut, "
    "renegotiate, or consolidate.\n"
    "Estimated impact: Prevents post-transformation cost drift. $200K to $300K annual savings by catching "
    "vendor creep that would otherwise erode the $2.62M structural savings.\n"
    "Complexity: Low. Runs on existing AP data within 4 to 6 weeks.\n\n"
    "90 Day AI Roadmap:\n"
    "  Week 1 to 4: I build the spend intelligence pipeline on existing AP data. Quick win.\n"
    "  Week 4 to 8: I pilot LLM reconciliation on the highest-volume entity.\n"
    "  Week 8 to 12: I roll out close automation across all entities. Target: sub-5 day close."
)

AI_TOOLS_USED = (
    "1. What AI tools I used\n\n"
    "Claude Code (Anthropic CLI agent) was the primary tool. "
    "I used it to parse the Input P&L workbook (8 sheets, 4,000+ line items), "
    "build the benchmark mapping, create the cost model against Central Finance roles, "
    "generate the 5 analysis sheets, and produce this DD1 document.\n\n"
    "Python (openpyxl, python-docx) handled file generation. Claude Code drove the analytical decisions: "
    "which benchmark category each line item maps to, how to structure the "
    "5 Why chain, and how to model the target team.\n\n"
    "2. How they helped my analysis\n\n"
    "Claude Code processed 2,092 OPEX rows and 1,000 COGS rows to build the benchmark mapping "
    "in minutes. It identified the F&A cost anomaly (4.8% of revenue for one "
    "sub-department exceeding the entire 4.5% Shared Services benchmark) by aggregating "
    "employee and non-employee data simultaneously. All 5 analysis sheets in the P&L workbook "
    "were generated programmatically from the 8 source tabs.\n\n"
    "The cost model was built by reading the Central Finance Roles spreadsheet and mapping current "
    "salaries to role tiers by salary band, producing the $2.62M savings figure with "
    "full traceability.\n\n"
    "3. How this approach translates to the role\n\n"
    "As VP of Operations, I run this methodology on every acquired business unit: "
    "ingest the P&L, map to benchmarks, identify the highest-gap function, deep dive to root cause, "
    "and model the Central Factory fix with specific roles and costs.\n\n"
    "The AI tooling makes this repeatable at acquisition pace. Instead of 4 to 6 weeks per BU, "
    "this approach produces a board-ready deep dive in days. That matters when the goal is one "
    "acquisition per week."
)


# ---------------------------------------------------------------------------
# Document generation
# ---------------------------------------------------------------------------

def generate_dd1() -> None:
    """Open template, fill all 8 rows, save output."""
    os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)

    doc = Document(TEMPLATE)
    table = doc.tables[0]

    # Row 0: Function - already "Operations"
    # Row 1: Playbook Item - already "Initial Import"

    # Row 2: Problem Statement
    _set_cell_text(table.rows[2].cells[1], PROBLEM_STATEMENT)

    # Row 3: 5 Why Analysis
    _set_cell_text(table.rows[3].cells[1], FIVE_WHY_ANALYSIS)

    # Row 4: Root Cause
    _set_cell_text(table.rows[4].cells[1], ROOT_CAUSE)

    # Row 5: Fix
    _set_cell_text(table.rows[5].cells[1], _build_fix_text())

    # Row 6: AI Opportunities
    _set_cell_text(table.rows[6].cells[1], AI_OPPORTUNITIES)

    # Row 7: AI Tools Used
    _set_cell_text(table.rows[7].cells[1], AI_TOOLS_USED)

    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")

    # Validate: check no empty cells
    doc2 = Document(OUTPUT)
    table2 = doc2.tables[0]
    for i in range(8):
        label = table2.rows[i].cells[0].text.strip()
        content = table2.rows[i].cells[1].text.strip()
        status = "OK" if content else "EMPTY"
        print(f"  Row {i} ({label}): {status} ({len(content)} chars)")


def _set_cell_text(cell, text: str) -> None:
    """Clear existing cell content and write new text with consistent formatting."""
    # Clear existing paragraphs
    for p in cell.paragraphs:
        p.clear()

    # Remove all but first paragraph
    while len(cell.paragraphs) > 1:
        p = cell.paragraphs[-1]
        p._element.getparent().remove(p._element)

    # Write text as formatted paragraphs
    lines = text.split("\n")
    first = True
    for line in lines:
        if first:
            para = cell.paragraphs[0]
            first = False
        else:
            para = cell.add_paragraph()

        run = para.add_run(line)
        run.font.name = "Arial"
        run.font.size = Pt(9)

        # Bold formatting for section headers
        if line.startswith("What:") or line.startswith("Why:"):
            run.bold = True
        elif line.startswith("Question"):
            run.bold = True
        elif line.startswith("Answer:") or line.startswith("Evidence:"):
            run.bold = False
        elif line.startswith("Area "):
            run.bold = True
        elif line.startswith("1.") or line.startswith("2.") or line.startswith("3."):
            if "What AI tools" in line or "How they helped" in line or "How this approach" in line:
                run.bold = True

        para.space_after = Pt(2)
        para.space_before = Pt(0)


if __name__ == "__main__":
    generate_dd1()
